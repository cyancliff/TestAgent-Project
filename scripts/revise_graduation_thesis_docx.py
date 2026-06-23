from __future__ import annotations

import os
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


def find_source_docx() -> Path:
    root = Path(os.environ["USERPROFILE"]) / "Desktop"
    candidates = [
        p
        for p in root.rglob("*.docx")
        if not p.name.startswith("~$")
        and p.name.startswith("张翼鹏毕业论文定稿")
        and "答辩前修订版" not in p.name
        and "扩充" not in p.name
    ]
    if not candidates:
        raise FileNotFoundError("未找到原始毕业论文 docx。")
    candidates.sort(key=lambda p: p.stat().st_size, reverse=True)
    return candidates[0]


def set_run_font(run, size: float | None = 10.5) -> None:
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")


def set_paragraph_font(paragraph: Paragraph, size: float | None = 10.5) -> None:
    for run in paragraph.runs:
        set_run_font(run, size)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def insert_picture_after(paragraph: Paragraph, image_path: Path, width_inches: float = 5.75) -> Paragraph:
    new_para = insert_paragraph_after(paragraph, "", paragraph.style)
    new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return new_para


def insert_paragraph_clone_after(paragraph: Paragraph, source_paragraph: Paragraph) -> Paragraph:
    cloned = deepcopy(source_paragraph._p)
    paragraph._p.addnext(cloned)
    return Paragraph(cloned, paragraph._parent)


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def find_para(doc: Document, predicate) -> Paragraph:
    for paragraph in doc.paragraphs:
        if predicate(paragraph):
            return paragraph
    raise ValueError("未找到目标段落。")


def find_para_index(doc: Document, predicate) -> int:
    for idx, paragraph in enumerate(doc.paragraphs):
        if predicate(paragraph):
            return idx
    raise ValueError("未找到目标段落索引。")


def find_last_para_index(doc: Document, predicate) -> int:
    found = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if predicate(paragraph):
            found = idx
    if found is None:
        raise ValueError("未找到目标段落索引。")
    return found


def delete_between(doc: Document, start_idx: int, end_idx: int) -> None:
    for paragraph in list(doc.paragraphs[start_idx:end_idx])[::-1]:
        delete_paragraph(paragraph)


def delete_xml_between(start_paragraph: Paragraph, end_paragraph: Paragraph) -> None:
    start = start_paragraph._p
    end = end_paragraph._p
    parent = start.getparent()
    node = start.getnext()
    while node is not None and node is not end:
        next_node = node.getnext()
        parent.remove(node)
        node = next_node


def style_or_fallback(doc: Document, style_name: str, fallback: str = "Normal"):
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles[fallback]


def add_sup_citation_after_snippet(doc: Document, snippet: str, citation: str) -> None:
    paragraph = find_para(doc, lambda p: snippet in p.text)
    if citation in paragraph.text:
        return
    run = paragraph.add_run(citation)
    run.font.superscript = True
    set_run_font(run)


def get_or_add_r_pr(r):
    r_pr = r.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r.insert(0, r_pr)
    return r_pr


def make_text_run(base_r, text: str, superscript: bool = False):
    new_r = OxmlElement("w:r")
    base_r_pr = base_r.find(qn("w:rPr"))
    if base_r_pr is not None:
        new_r.append(deepcopy(base_r_pr))
    if superscript:
        r_pr = get_or_add_r_pr(new_r)
        for old in list(r_pr.findall(qn("w:vertAlign"))):
            r_pr.remove(old)
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), "superscript")
        r_pr.append(vert)
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    new_r.append(t)
    return new_r


CITATION_RE = re.compile(r"\[(\d+(?:\s*-\s*\d+)?(?:\s*,\s*\d+(?:\s*-\s*\d+)?)*)\]")


def citation_numbers(text: str) -> list[int]:
    nums: list[int] = []
    for part in re.split(r"\s*,\s*", text):
        if "-" in part:
            left, right = [int(x) for x in re.split(r"\s*-\s*", part)]
            nums.extend(range(left, right + 1))
        else:
            nums.append(int(part))
    return nums


def is_reference_citation(match: re.Match) -> bool:
    nums = citation_numbers(match.group(1))
    return bool(nums) and all(1 <= n <= 43 for n in nums)


def superscript_body_citations(doc: Document) -> None:
    ref_idx = find_para_index(doc, lambda p: p.text.strip() == "参考文献")
    for paragraph in doc.paragraphs[:ref_idx]:
        if not CITATION_RE.search(paragraph.text):
            continue
        for run in list(paragraph.runs):
            text = run.text
            matches = [m for m in CITATION_RE.finditer(text) if is_reference_citation(m)]
            if not matches:
                continue
            parent = run._r.getparent()
            insert_at = parent.index(run._r)
            cursor = 0
            new_runs = []
            for match in matches:
                if match.start() > cursor:
                    new_runs.append(make_text_run(run._r, text[cursor : match.start()], False))
                new_runs.append(make_text_run(run._r, match.group(0), True))
                cursor = match.end()
            if cursor < len(text):
                new_runs.append(make_text_run(run._r, text[cursor:], False))
            for new_r in new_runs:
                parent.insert(insert_at, new_r)
                insert_at += 1
            parent.remove(run._r)


def replace_algorithm_blocks(doc: Document) -> None:
    body_style = style_or_fallback(doc, "标准正文")

    alg1_start = find_para_index(doc, lambda p: p.text.startswith("算法 3.1 ATMR-CAT"))
    alg1_end = find_para_index(doc, lambda p: p.text.startswith("综上，ATMR-CAT"))
    alg1_anchor = doc.paragraphs[alg1_start - 1]
    delete_between(doc, alg1_start, alg1_end)
    alg1_lines = [
        "算法 3.1 ATMR-CAT 自适应选题算法",
        "算法思想：该算法在固定 ATMR 阶段题库内工作，不生成新题，也不改变题干和选项。每轮先合并正式作答与阶段内临时作答，再依据作答得分、题目区分度和异常降权因子更新后验估计，并在当前阶段候选题中综合信息量、覆盖度、难度匹配和区分度选择下一题。",
        "输入：当前阶段 m，正式作答 Asaved，阶段内临时作答 Atemp，阶段题库 Qm，题目参数 aq、bq、vq，当前后验 μt、σt2，作答得分 si，异常降权因子 ηi，阶段最大题量 Nm。",
        "输出：下一题 q* 或阶段结束标记 stop，更新后的后验 μt+1、σt+12，候选题评分表 R。",
        "1. 合并作答 Am ← Asaved ∪ Atemp；若 |Am| ≥ Nm，则返回 stop。",
        "2. 若 Am = ∅，在当前阶段启用题目中按 Sfirst(q)=0.6(1-2|bq-0.5|)+0.4aq 选择初始题 q* 并返回。",
        "3. 对 Am 中每个已答题 i，令 zi=si/5、ρi=2aiηi，逐题更新后验精度、均值和方差，得到 μt+1 与 σt+12。",
        "4. 构造候选集 Qcand={q | q∈Qm，q is enabled，q∉Am}；若候选集为空，则返回 stop。",
        "5. 根据不确定性 u=min(σt2/0.25,1) 更新 λI、λC、λD、λA。",
        "6. 对每个 q∈Qcand 计算 I(q)、C(q)、D(q)、A(q)，并得到 S(q)=λII(q)+λCC(q)+λDD(q)+λAA(q)。",
        "7. 若候选评分无效，则按阶段顺序回退选择首个可用题；否则 q*=argmax S(q)。",
        "8. 返回 q*、μt+1、σt+12 和 R。",
    ]
    anchor = alg1_anchor
    for line in alg1_lines:
        anchor = insert_paragraph_after(anchor, line, body_style)

    alg2_start = find_para_index(doc, lambda p: p.text.startswith("算法 3.2 异常作答检测"))
    alg2_end = find_para_index(doc, lambda p: p.text.startswith("异常作答检测与可信度评估流程"))
    alg2_anchor = doc.paragraphs[alg2_start - 1]
    delete_between(doc, alg2_start, alg2_end)
    alg2_lines = [
        "算法 3.2 异常作答检测与可信度评估算法",
        "算法思想：该算法将前端压缩后的作答行为摘要转化为题级风险分，再由题级风险得到异常作答标记和单题可信度，并按 ATMR 维度汇总为维度可信度、整体可信度和可信度加权参考分。算法只描述作答证据质量，不判断用户主观动机，也不替代原始 ATMR 维度得分。",
        "输入：当前会话作答集合 A，题级行为摘要 Hi，风险规则集合 Rrule，规则风险增量 Δik，题目得分 si，维度 m 的已作答题目集合 Sm，期望题量 Nm。",
        "输出：题级风险分 ri，异常作答标记 zi，单题可信度 ci，维度可信度 Cm，整体可信度 C，可信度加权参考分 scoremconf。",
        "1. 对每个已答题 i，遍历风险规则 Rrule，根据 Hi 累加被触发规则的 Δik，得到 ri=min(ΣkΔik,100)。",
        "2. 若 ri≥50，则 zi=1，否则 zi=0；计算 ci=clamp(1-0.0065ri-0.08zi,0.2,1.0)，并保存风险原因。",
        "3. 对每个维度 m，取 Sm，计算 nm=|Sm|、am=Σi∈Sm zi 与 completionm=min(nm/Nm,1)。",
        "4. 计算 Cm=clamp(meani∈Sm(ci)·(0.85+0.15·completionm)-0.015am,0,1)。",
        "5. 计算 scoremconf=(Σi∈Sm ci si/Σi∈Sm ci)·nm，用于与原始维度得分进行稳健性比较。",
        "6. 计算整体异常数 Aabn=Σi zi，并得到 C=clamp(mean(ci)-min(0.01·Aabn,0.12),0,1)。",
        "7. 将 Cm 与 C 映射为“较高”“中等”“较低”标签，返回 ri、zi、ci、Cm、C 和 scoremconf。",
    ]
    anchor = alg2_anchor
    for line in alg2_lines:
        anchor = insert_paragraph_after(anchor, line, body_style)


def replace_chapter5_run_flow(doc: Document, resource_dir: Path) -> None:
    body_style = style_or_fallback(doc, "标准正文")
    caption_style = style_or_fallback(doc, "表格")

    heading_idx = find_last_para_index(doc, lambda p: p.text.startswith("5.4.3 系统运行流程"))
    next_heading_idx = find_last_para_index(doc, lambda p: p.text.startswith("5.4.4 系统功能测试"))
    atmr_report_image = doc.paragraphs[
        find_last_para_index(doc, lambda p: p.text.strip() == "图5-7 ATMR 报告总览与四维人格画像页面") - 1
    ]
    ai_consult_image = doc.paragraphs[
        find_last_para_index(doc, lambda p: p.text.strip() == "图5-8 关联 ATMR 报告的智能咨询页面") - 1
    ]
    big_five_report_image = doc.paragraphs[
        find_last_para_index(doc, lambda p: p.text.strip() == "图5-10 大五人格预测报告详情页面") - 1
    ]
    cloned_atmr_report = deepcopy(atmr_report_image._p)
    cloned_ai_consult = deepcopy(ai_consult_image._p)
    cloned_big_five_report = deepcopy(big_five_report_image._p)
    anchor = doc.paragraphs[heading_idx]
    delete_xml_between(anchor, doc.paragraphs[next_heading_idx])

    content = [
        (
            "p",
            "系统运行流程由前端交互、后端业务服务、智能分析任务和结果展示共同完成。为避免第五章停留在页面展示层面，本文将运行过程拆分为系统实现结构、ATMR 测评状态流、智能报告生成流程和多模态异步任务流四类结构说明，并保留 ATMR 报告、大五报告和 AI 咨询三类关键页面作为用户侧结果验证。",
        ),
        ("img", resource_dir / "fig5_backend_modules.png"),
        ("cap", "图5-2 系统实现结构图"),
        (
            "p",
            "如图5-2所示，后端以 FastAPI 接口层承接用户请求，并通过用户权限、测评会话、报告生成、多模态任务和咨询会话等服务模块组织业务逻辑。数据库层保存用户、作答记录、报告、视频任务和聊天消息，智能分析层则封装 ATMR-CAT、可信度评估、PageIndex 检索、多智能体报告生成与多模态推理能力。各模块通过用户归属关系和任务状态衔接，避免在业务层混合不同来源证据。",
        ),
        ("img", resource_dir / "fig5_atmr_state_flow.png"),
        ("cap", "图5-3 ATMR测评与报告生成状态流"),
        (
            "p",
            "ATMR 主流程如图5-3所示。用户进入测评后，系统优先检查未完成会话并恢复阶段状态；每次提交阶段作答时，后台完成题目归属校验、作答记录写入、题级可信度计算和阶段状态推进。A、T、M、R 四个阶段完成后，系统读取模块层分析结果并触发综合报告生成，最终将报告保存到用户历史记录中。",
        ),
        ("img", resource_dir / "fig5_report_generation_flow.png"),
        ("cap", "图5-4 智能报告生成流程图"),
        (
            "p",
            "智能报告生成流程如图5-4所示。阶段作答记录、维度得分、题目证据和可信度摘要先被整理为结构化测评证据，随后由 PageIndex 检索对应 ATMR 知识材料。模块层按优势分析、风险审查和综合裁决形成单维度总结，综合层再汇总四个维度的裁决结果，生成整体画像、发展建议和解释边界说明。该流程使大模型输出始终受到作答证据、知识材料和可信度信息共同约束。",
        ),
        ("img", resource_dir / "fig5_multimodal_async.png"),
        ("cap", "图5-5 多模态人格分析异步任务流程"),
        (
            "p",
            "多模态人格分析流程如图5-5所示。用户上传短视频后，系统先创建大五人格报告任务并返回可见状态，后台异步执行视频抽帧、音频提取、语音转写、视觉/文本/音频特征提取、背景语义补充和可选 MOL 微表情特征提取。特征包校验通过后进入 AGTN-MTL checkpoint 推理；若依赖缺失或推理失败，则记录失败或回退状态，避免以占位结果替代真实模型输出。",
        ),
        ("clone_img", cloned_atmr_report),
        ("cap", "图5-6 ATMR报告总览关键页面"),
        (
            "p",
            "如图5-6所示，ATMR 报告总览页面用于说明系统最终向用户呈现维度得分、可信度摘要、报告正文和题目证据入口。该页面对应 ATMR 主链路的最终结果展示，能够验证前述状态流、报告生成和证据追溯设计是否落到用户可见结果中。",
        ),
        ("clone_img", cloned_big_five_report),
        ("cap", "图5-7 大五人格报告详情关键页面"),
        (
            "p",
            "如图5-7所示，大五人格报告详情页面用于说明视频人格分析辅线的结果呈现方式。页面展示五维人格预测分数、任务状态、模型版本、模态质量和智能解读内容，与图5-5中的异步任务流相对应，同时强调该结果独立于 ATMR 问卷计分。",
        ),
        ("clone_img", cloned_ai_consult),
        ("cap", "图5-8 关联报告的 AI 咨询关键页面"),
        (
            "p",
            "如图5-8所示，关联报告的 AI 咨询页面用于说明系统如何在咨询场景中读取用户显式选择的报告上下文。咨询模块可以参考 ATMR 报告和大五报告，但回答时需要区分问卷自陈证据与视频外显线索，避免将不同证据链合并为单一诊断结论。",
        ),
        (
            "p",
            "综上，系统运行流程覆盖用户认证、阶段化作答、报告生成、报告后咨询、视频人格分析和后台质量治理等关键环节。与原稿逐页展示登录、注册、工作台和管理端截图相比，本节保留三类具有结果验证意义的关键页面，并用结构图说明数据流、状态流和异步任务组织方式，使第五章更侧重系统实现而不是软件说明书式页面介绍。",
        ),
    ]

    for kind, payload in content:
        if kind == "p":
            anchor = insert_paragraph_after(anchor, payload, body_style)
        elif kind == "img":
            anchor = insert_picture_after(anchor, payload)
        elif kind == "clone_img":
            cloned = OxmlElement("w:p")
            cloned = deepcopy(payload)
            anchor._p.addnext(cloned)
            anchor = Paragraph(cloned, anchor._parent)
            anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "cap":
            anchor = insert_paragraph_after(anchor, payload, caption_style)
            anchor.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_font(anchor)


REFERENCES = [
    "[1] CHEN D, LIU Y, GUO Y, et al. The revolution of generative artificial intelligence in psychology: The interweaving of behavior, consciousness, and ethics[J]. Acta Psychologica, 2024, 251: 104593.",
    "[2] HONG Y, XIA Z. AI-driven innovations in psychological assessment: Multimodal data, intelligent analytics, and ethical challenges[C]//Proceedings of the 2025 International Conference on Artificial Intelligence and Smart Manufacturing. 2025: 854-859.",
    "[3] WEISS D J. Improving measurement quality and efficiency with adaptive testing[J]. Applied Psychological Measurement, 1982, 6(4): 473-492.",
    "[4] YU J, ZHUANG Y, SUN Y, et al. TestAgent: An adaptive and intelligent expert for human assessment[C]//Findings of the Association for Computational Linguistics: ACL 2025. Vienna: Association for Computational Linguistics, 2025: 724-747.",
    "[5] LI X, CHEN X, NIU Y, et al. Psydi: Towards a personalized and progressively in-depth chatbot for psychological measurements[J/OL]. arXiv:2408.03337, 2024.",
    "[6] MILANO N, PONTICORVO M, MAROCCO D. Human expertise and large language model embeddings in the content validity assessment of personality tests[J]. Educational and Psychological Measurement, 2025: 00131644251355485.",
    "[7] RAVENDA F, PRETI A, POLETTI M, et al. Rethinking psychometrics through LLMs: How item semantics shape measurement and prediction in psychological questionnaires[J]. Scientific Reports, 2025, 15(1): 37313.",
    "[8] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-augmented generation for knowledge-intensive NLP tasks[C]//Advances in Neural Information Processing Systems. 2020, 33: 9459-9474.",
    "[9] ASAI A, WU Z, WANG Y, et al. Self-RAG: Learning to retrieve, generate, and critique through self-reflection[J/OL]. arXiv:2310.11511, 2024.",
    "[10] HU W, ZHANG W, JIANG Y, et al. Removal of hallucination on hallucination: Debate-augmented RAG[C]//Proceedings of the 63rd Annual Meeting of the Association for Computational Linguistics. 2025: 15839-15853.",
    "[11] DU Y, LI S, TORRALBA A, et al. Improving factuality and reasoning in language models through multiagent debate[C]//Forty-first International Conference on Machine Learning. 2024.",
    "[12] JOHN O P, SRIVASTAVA S. The Big-Five trait taxonomy: History, measurement, and theoretical perspectives[M]//PERVIN L A, JOHN O P. Handbook of Personality: Theory and Research. 2nd ed. New York: Guilford Press, 1999: 102-138.",
    "[13] COSTA P T, MCCRAE R R. Revised NEO Personality Inventory and NEO Five-Factor Inventory Professional Manual[M]. Odessa: Psychological Assessment Resources, 1992.",
    "[14] PONCE-LÓPEZ V, CHEN B, OLIU M, et al. ChaLearn LAP 2016: First round challenge on first impressions-dataset and results[C]//European Conference on Computer Vision Workshops. 2016: 400-418.",
    "[15] GÜÇLÜTÜRK Y, GÜÇLÜ U, VAN GERVEN M A J, et al. Deep impression: Audiovisual deep residual networks for multimodal apparent personality trait recognition[C]//European Conference on Computer Vision Workshops. 2016: 349-358.",
    "[16] ASLAN S, GÜDÜKBAY U. Multimodal video-based apparent personality recognition using long short-term memory and convolutional neural networks[EB/OL]. arXiv:1911.00381, 2019.",
    "[17] RADFORD A, KIM J W, HALLACY C, et al. Learning transferable visual models from natural language supervision[C]//International Conference on Machine Learning. 2021.",
    "[18] WU H H, SEETHARAMAN P, KUMAR K, et al. Wav2CLIP: Learning robust audio representations from CLIP[C]//IEEE International Conference on Acoustics, Speech and Signal Processing. 2022.",
    "[19] VASWANI A, SHAZEER N, PARMAR N, et al. Attention is all you need[C]//Advances in Neural Information Processing Systems. 2017.",
    "[20] WANG R, ZHAO X, XU X, et al. A multimodal personality prediction framework based on adaptive graph transformer network and multi-task learning[J]. Computer Graphics Forum, 2025, 44(2): e70030.",
    "[21] ZOU H, WANG P, YAN Z, et al. Can LLM self-report?: Evaluating the validity of self-report scales in measuring personality design in LLM-based chatbots[J/OL]. arXiv:2412.00207, 2024.",
    "[22] LIU Y F, LU Y L, HE D, et al. From Five Dimensions to Many: Large Language Models as Precise and Interpretable Psychological Profilers[J/OL]. arXiv:2511.03235, 2025.",
    "[23] ZHANG M, TANG Y, PAGEINDEX TEAM. PageIndex: Next-Generation Vectorless, Reasoning-based RAG[EB/OL]. (2025-09-19)[2026-05-30]. https://pageindex.ai/blog/pageindex-intro.",
    "[24] OH Y H, SEE J, LE NGO A C, et al. A survey of automatic facial micro-expression analysis: Databases, methods, and challenges[J]. Frontiers in Psychology, 2018, 9: 1128.",
    "[25] LI X, HONG X, MOILANEN A, et al. Towards reading hidden emotions: A comparative study of spontaneous micro-expression spotting and recognition methods[J]. IEEE Transactions on Affective Computing, 2018, 9(4): 563-577.",
    "[26] SHAO Z, CHENG Y, LI F, et al. MOL: Joint estimation of micro-expression, optical flow, and landmark via transformer-graph-style convolution[J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2025, 47(10): 8756-8768.",
    "[27] WAINER H, DORANS N J, FLAUGHER R, et al. Computerized Adaptive Testing: A Primer[M]. 2nd ed. Mahwah: Lawrence Erlbaum Associates, 2000.",
    "[28] SAMEJIMA F. Estimation of latent ability using a response pattern of graded scores[J]. Psychometrika Monograph Supplement, 1969, 34(4): 1-100.",
    "[29] CRONBACH L J. Coefficient alpha and the internal structure of tests[J]. Psychometrika, 1951, 16(3): 297-334.",
    "[30] 戴海崎, 张锋, 陈雪枫. 心理与教育测量[M]. 4版. 广州: 暨南大学出版社, 2018.",
    "[31] 郑日昌, 吴九君. 心理测量学[M]. 北京: 人民教育出版社, 1999.",
    "[32] GAO Y, XIONG Y, GAO X, et al. Retrieval-augmented generation for large language models: A survey[J/OL]. arXiv:2312.10997, 2023.",
    "[33] JI Z, LEE N, FRIESKE R, et al. Survey of hallucination in natural language generation[J]. ACM Computing Surveys, 2023, 55(12): 1-38.",
    "[34] REIMERS N, GUREVYCH I. Sentence-BERT: Sentence embeddings using Siamese BERT-networks[C]//Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing and the 9th International Joint Conference on Natural Language Processing. 2019: 3982-3992.",
    "[35] RADFORD A, KIM J W, XU T, et al. Robust speech recognition via large-scale weak supervision[C]//Proceedings of the 40th International Conference on Machine Learning. 2023: 28492-28518.",
    "[36] 国家互联网信息办公室, 中华人民共和国国家发展和改革委员会, 中华人民共和国教育部, 等. 生成式人工智能服务管理暂行办法[EB/OL]. (2023-07-13)[2026-06-18]. https://www.cac.gov.cn/2023-07/13/c_1690898327029107.htm.",
    "[37] 全国人民代表大会常务委员会. 中华人民共和国个人信息保护法[EB/OL]. (2021-08-20)[2026-06-18]. http://www.npc.gov.cn/npc/c30834/202108/t20210820_313088.html.",
    "[38] YAN W J, LI X, WANG S J, et al. CASME II: An improved spontaneous micro-expression database and the baseline evaluation[J]. PLoS ONE, 2014, 9(1): e86041.",
    "[39] LI X, PFISTER T, HUANG X, et al. A spontaneous micro-expression database: Inducement, collection and baseline[C]//IEEE International Conference and Workshops on Automatic Face and Gesture Recognition. 2013: 1-6.",
    "[40] DAVISON A K, LANSLEY C, COSTEN N, et al. SAMM: A spontaneous micro-facial movement dataset[J]. IEEE Transactions on Affective Computing, 2018, 9(1): 116-129.",
    "[41] 中华人民共和国国家市场监督管理总局, 中国国家标准化管理委员会. 信息安全技术 个人信息安全规范: GB/T 35273-2020[S]. 北京: 中国标准出版社, 2020.",
    "[42] 中华人民共和国国家市场监督管理总局, 中国国家标准化管理委员会. 信息技术 人工智能 术语: GB/T 41867-2022[S]. 北京: 中国标准出版社, 2022.",
    "[43] 科学技术部, 教育部, 工业和信息化部, 等. 科技伦理审查办法(试行)[EB/OL]. (2023-09-07)[2026-06-18]. https://www.most.gov.cn/xxgk/xinxifenlei/fdzdgknr/fgzc/gfxwj/gfxwj2023/202310/t20231008_188064.html.",
]


def replace_references(doc: Document) -> None:
    ref_idx = find_last_para_index(doc, lambda p: p.text.strip() == "参考文献")
    appendix_idx = find_last_para_index(doc, lambda p: p.text.startswith("附录A"))
    ref_title = doc.paragraphs[ref_idx]
    ref_title.paragraph_format.page_break_before = True
    ref_style = style_or_fallback(doc, "参考文献正文")
    delete_xml_between(ref_title, doc.paragraphs[appendix_idx])
    anchor = ref_title
    for entry in REFERENCES:
        anchor = insert_paragraph_after(anchor, entry, ref_style)
        set_paragraph_font(anchor)
        anchor.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def polish_chapter5_summary(doc: Document) -> None:
    old = "通过截图界面和测试任务说明关键流程如何落到用户界面和后台状态中"
    new = "通过实现结构图、关键页面和测试任务说明关键流程如何落到系统模块、用户界面和后台状态中"
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    set_run_font(run)
                    return
            paragraph.text = paragraph.text.replace(old, new)
            set_paragraph_font(paragraph)
            return


def remove_page_break_before_chapter6(doc: Document) -> None:
    chapter6_idx = find_last_para_index(doc, lambda p: p.text.startswith("第六章 总结与展望"))
    if chapter6_idx == 0:
        return
    prev = doc.paragraphs[chapter6_idx - 1]
    if prev.text.strip() == "" and prev._p.xpath(".//w:br[@w:type='page']"):
        delete_paragraph(prev)


def set_all_table_fonts(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, 10.5)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    for old in list(settings.findall(qn("w:updateFields"))):
        settings.remove(old)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def load_diagram_font(size: int, bold: bool = False):
    font_names = [
        "msyhbd.ttc" if bold else "msyh.ttc",
        "simhei.ttf",
        "simsun.ttc",
    ]
    for name in font_names:
        path = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / name
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


DIAGRAM_W = 2400
DIAGRAM_H = 1350
INK = (31, 42, 55)
MUTED = (86, 102, 121)
LINE = (91, 118, 151)
GRID = (214, 222, 231)
BG = (248, 250, 252)
PANEL = (255, 255, 255)
BLUE = ((226, 239, 252), (43, 105, 173))
TEAL = ((226, 244, 239), (45, 132, 113))
AMBER = ((255, 244, 224), (181, 116, 37))
VIOLET = ((241, 233, 249), (126, 87, 170))
ROSE = ((255, 238, 238), (190, 81, 81))
SLATE = ((238, 242, 247), (92, 110, 132))


def wrapped_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        trial = current + char
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def text_size(draw: ImageDraw.ImageDraw, text: str, font) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str) -> None:
    title_font = load_diagram_font(72, True)
    sub_font = load_diagram_font(34)
    draw.text((96, 62), title, font=title_font, fill=INK)
    draw.text((100, 148), subtitle, font=sub_font, fill=MUTED)
    draw.line((96, 214, DIAGRAM_W - 96, 214), fill=GRID, width=3)


def draw_badge(draw: ImageDraw.ImageDraw, box, text: str, fill, outline) -> None:
    x1, y1, x2, y2 = box
    font = load_diagram_font(30, True)
    draw.rounded_rectangle(box, radius=24, fill=fill, outline=outline, width=3)
    tw, th = text_size(draw, text, font)
    draw.text((x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2 - 2), text, font=font, fill=outline)


def draw_round_box(draw: ImageDraw.ImageDraw, box, title: str, desc: str, fill, outline, note: str | None = None) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle((x1 + 7, y1 + 9, x2 + 7, y2 + 9), radius=28, fill=(226, 232, 240))
    draw.rounded_rectangle(box, radius=28, fill=fill, outline=outline, width=4)
    title_font = load_diagram_font(44, True)
    desc_font = load_diagram_font(32)
    note_font = load_diagram_font(28)
    draw.text((x1 + 34, y1 + 30), title, font=title_font, fill=INK)
    y = y1 + 92
    for line in wrapped_text(draw, desc, desc_font, x2 - x1 - 68):
        draw.text((x1 + 34, y), line, font=desc_font, fill=(55, 65, 81))
        y += 42
    if note:
        draw.line((x1 + 34, y2 - 68, x2 - 34, y2 - 68), fill=(203, 213, 225), width=2)
        draw.text((x1 + 34, y2 - 50), note, font=note_font, fill=MUTED)


def draw_arrow(draw: ImageDraw.ImageDraw, start, end, color=LINE, width: int = 7) -> None:
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        points = [(ex, ey), (ex - 26 * direction, ey - 15), (ex - 26 * direction, ey + 15)]
    else:
        direction = 1 if ey >= sy else -1
        points = [(ex, ey), (ex - 15, ey - 26 * direction), (ex + 15, ey - 26 * direction)]
    draw.polygon(points, fill=color)


def create_diagram_canvas(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (DIAGRAM_W, DIAGRAM_H), BG)
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((44, 38, DIAGRAM_W - 44, DIAGRAM_H - 38), radius=34, fill=PANEL, outline=(220, 226, 235), width=3)
    draw_title(draw, title, subtitle)
    return img, draw


def draw_labeled_arrow(draw: ImageDraw.ImageDraw, start, end, label: str | None = None, color=LINE) -> None:
    draw_arrow(draw, start, end, color=color)
    if label:
        font = load_diagram_font(28)
        sx, sy = start
        ex, ey = end
        mx, my = (sx + ex) / 2, (sy + ey) / 2
        tw, th = text_size(draw, label, font)
        pad = 12
        draw.rounded_rectangle((mx - tw / 2 - pad, my - th / 2 - pad, mx + tw / 2 + pad, my + th / 2 + pad), radius=16, fill=(255, 255, 255), outline=(203, 213, 225), width=2)
        draw.text((mx - tw / 2, my - th / 2 - 1), label, font=font, fill=MUTED)


def draw_footer_note(draw: ImageDraw.ImageDraw, text: str) -> None:
    font = load_diagram_font(30)
    x1, y1, x2, y2 = 110, 1190, DIAGRAM_W - 110, 1270
    draw.rounded_rectangle((x1, y1, x2, y2), radius=24, fill=(248, 250, 252), outline=(215, 224, 235), width=2)
    draw.text((x1 + 32, y1 + 23), text, font=font, fill=(71, 85, 105))


def ensure_backend_modules_diagram(resource_dir: Path) -> None:
    path = resource_dir / "fig5_backend_modules.png"
    img, draw = create_diagram_canvas("系统实现结构", "FastAPI 服务层、数据持久层与智能分析能力的模块边界")
    boxes = [
        ((120, 330, 510, 530), "用户与权限", "身份认证\n资源归属校验", BLUE),
        ((710, 330, 1100, 530), "ATMR测评", "会话恢复\n阶段作答", TEAL),
        ((1300, 330, 1690, 530), "智能报告", "知识检索\n多智能体审议", AMBER),
        ((1890, 330, 2280, 530), "多模态分析", "视频任务\n大五报告", VIOLET),
        ((430, 770, 820, 970), "数据持久层", "用户、作答\n报告、消息", SLATE),
        ((1030, 770, 1420, 970), "知识与任务层", "ATMR知识库\n异步任务状态", BLUE),
        ((1630, 770, 2020, 970), "咨询与治理", "报告上下文\n质量追踪", TEAL),
    ]
    for box, title, desc, palette in boxes:
        draw_round_box(draw, box, title, desc, palette[0], palette[1])
    for s, e, label in [
        ((510, 420), (710, 420), "会话"),
        ((1100, 420), (1300, 420), "证据"),
        ((1690, 420), (1890, 420), "辅线"),
        ((905, 520), (665, 760), "写入"),
        ((1495, 520), (1225, 760), "检索"),
        ((2085, 520), (1825, 760), "状态"),
        ((1420, 860), (1630, 860), "上下文"),
    ]:
        draw_labeled_arrow(draw, s, e, label)
    draw_footer_note(draw, "实现重点：按用户归属和任务状态解耦，ATMR主链路与视频辅链路并列保存、独立解释。")
    img.save(path)


def ensure_atmr_state_diagram(resource_dir: Path) -> None:
    path = resource_dir / "fig5_atmr_state_flow.png"
    img, draw = create_diagram_canvas("ATMR测评与报告生成状态流", "阶段作答、质量检测、模块审议与综合报告的状态推进")
    y = 400
    w, h = 300, 180
    xs = [95, 445, 795, 1145, 1495, 1845]
    nodes = [
        ("创建/恢复会话", "检查未完成记录\n确定当前阶段", BLUE),
        ("阶段作答", "按A/T/M/R推进\n保存答案", TEAL),
        ("质量检测", "行为摘要\n题级可信度", AMBER),
        ("模块审议", "优势、风险、裁决\n形成维度总结", VIOLET),
        ("综合报告", "汇总四维结果\n生成最终报告", BLUE),
        ("咨询引用", "用户显式选择\n构造上下文", TEAL),
    ]
    for x, (title, desc, palette) in zip(xs, nodes):
        draw_round_box(draw, (x, y, x + w, y + h), title, desc, palette[0], palette[1])
    for i in range(len(xs) - 1):
        draw_labeled_arrow(draw, (xs[i] + w, y + h // 2), (xs[i + 1], y + h // 2))
    bottom = [
        ((430, 790, 800, 965), "作答记录", "得分、耗时\n异常原因", SLATE),
        ((1110, 790, 1510, 965), "模块结果", "维度总结\n证据边界", SLATE),
        ((1580, 790, 1980, 965), "测评报告", "报告正文\n可信度摘要", SLATE),
    ]
    for box, title, desc, palette in bottom:
        draw_round_box(draw, box, title, desc, palette[0], palette[1])
    draw_labeled_arrow(draw, (595, 580), (615, 790), "落库", color=(100, 116, 139))
    draw_labeled_arrow(draw, (1295, 580), (1310, 790), "保存", color=(100, 116, 139))
    draw_labeled_arrow(draw, (1645, 580), (1780, 790), "持久化", color=(100, 116, 139))
    draw_footer_note(draw, "实现重点：阶段状态只推进当前维度；可信度不拦截作答，而是在报告解释中控制强度。")
    img.save(path)


def ensure_report_flow_diagram(resource_dir: Path) -> None:
    path = resource_dir / "fig5_report_generation_flow.png"
    img, draw = create_diagram_canvas("智能报告生成流程", "结构化测评证据、PageIndex RAG 与多智能体分层分析共同约束输出")
    top_y = 330
    boxes = [
        ((110, top_y, 500, top_y + 200), "证据整理", "维度得分\n题目证据与可信度", BLUE),
        ((690, top_y, 1080, top_y + 200), "知识检索", "召回ATMR章节\n补充解释边界", TEAL),
        ((1270, top_y, 1660, top_y + 200), "模块审议", "优势、风险、裁决\n形成维度总结", AMBER),
        ((1850, top_y, 2240, top_y + 200), "综合报告", "整体画像\n建议与边界", VIOLET),
        ((390, 770, 790, 970), "边界控制", "不诊断\n低可信降语气", ROSE),
        ((1000, 770, 1400, 970), "报告落库", "正文、证据摘要\n生成状态", BLUE),
        ((1610, 770, 2010, 970), "咨询上下文", "用户选择报告\n区分证据来源", TEAL),
    ]
    for box, title, desc, palette in boxes:
        draw_round_box(draw, box, title, desc, palette[0], palette[1])
    draw_labeled_arrow(draw, (500, top_y + 100), (690, top_y + 100))
    draw_labeled_arrow(draw, (1080, top_y + 100), (1270, top_y + 100))
    draw_labeled_arrow(draw, (1660, top_y + 100), (1850, top_y + 100))
    draw_labeled_arrow(draw, (2050, top_y + 200), (1810, 770), "进入咨询")
    draw_labeled_arrow(draw, (1850, 870), (1400, 870), None)
    draw_labeled_arrow(draw, (1000, 870), (790, 870), None)
    draw_labeled_arrow(draw, (1400, 870), (1610, 870), None)
    draw_footer_note(draw, "实现重点：模块层处理单维度证据，综合层汇总；报告和咨询均保留证据来源与可信度边界。")
    img.save(path)


def ensure_multimodal_async_diagram(resource_dir: Path) -> None:
    path = resource_dir / "fig5_multimodal_async.png"
    img, draw = create_diagram_canvas("多模态人格分析异步任务流程", "用户上传视频后，后台任务完成特征提取、模型推理与结果边界控制")
    y = 360
    w, h = 310, 180
    xs = [90, 470, 850, 1230, 1610, 1990]
    nodes = [
        ("上传视频", "创建任务\n返回处理中状态", BLUE),
        ("预处理", "抽帧、提取音频\n语音转写", TEAL),
        ("特征提取", "视觉、文本、音频\n统一抽取", AMBER),
        ("特征包校验", "长度、维度\n标签顺序", VIOLET),
        ("模型推理", "加载模型权重\n输出五维分数", BLUE),
        ("报告展示", "分数、模态质量\n独立解释", TEAL),
    ]
    for x, (title, desc, palette) in zip(xs, nodes):
        draw_round_box(draw, (x, y, x + w, y + h), title, desc, palette[0], palette[1])
    for i in range(len(xs) - 1):
        draw_labeled_arrow(draw, (xs[i] + w, y + h // 2), (xs[i + 1], y + h // 2))
    bottom = [
        ((470, 780, 820, 960), "失败状态", "依赖缺失\n记录错误原因", ROSE),
        ((1170, 780, 1520, 960), "可选MOL线索", "短时面部动态\n仅作补充特征", AMBER),
        ((1870, 780, 2220, 960), "结果边界", "不能诊断筛选\n不替代ATMR", ROSE),
    ]
    for box, title, desc, palette in bottom:
        draw_round_box(draw, box, title, desc, palette[0], palette[1])
    draw_labeled_arrow(draw, (625, y + h), (645, 780), "异常")
    draw_labeled_arrow(draw, (1385, y + h), (1345, 780), "补充")
    draw_labeled_arrow(draw, (2145, y + h), (2045, 780), "约束")
    draw_footer_note(draw, "实现重点：视频分析以任务状态驱动，失败不生成占位结论；大五报告作为独立辅线进入咨询。")
    img.save(path)


def ensure_all_diagrams(resource_dir: Path) -> None:
    ensure_backend_modules_diagram(resource_dir)
    ensure_atmr_state_diagram(resource_dir)
    ensure_report_flow_diagram(resource_dir)
    ensure_multimodal_async_diagram(resource_dir)


def apply_new_citations(doc: Document) -> None:
    additions = [
        ("真实样本再校准的候选对象。", "[27-30]"),
        ("ATMR-CAT 不仅改善潜变量估计，也能够提高题目证据多样性。", "[31]"),
        ("不能替代用户作答证据或专业心理评估。", "[32,33]"),
        ("不将视频结果写成 ATMR 计分依据。", "[34]"),
        ("完成抽帧、音频提取和语音转写", "[35]"),
        ("不能作为临床诊断、筛选或决策依据。", "[36,37]"),
        ("SAMM、CASME2 和 SMIC 三套帧序列样本", "[38-40]"),
        ("报告、视频和聊天记录需执行用户权限校验", "[41,42]"),
        ("访问控制、日志脱敏、文件删除和权限隔离", "[43]"),
    ]
    for snippet, citation in additions:
        add_sup_citation_after_snippet(doc, snippet, citation)


def main() -> None:
    src = find_source_docx()
    out = src.with_name(src.stem + "_答辩前修订版.docx")
    resource_dir = src.parent / "答辩前修订资源"
    resource_dir.mkdir(exist_ok=True)
    ensure_all_diagrams(resource_dir)

    doc = Document(src)
    apply_new_citations(doc)
    replace_algorithm_blocks(doc)
    replace_chapter5_run_flow(doc, resource_dir)
    polish_chapter5_summary(doc)
    remove_page_break_before_chapter6(doc)
    replace_references(doc)
    superscript_body_citations(doc)
    set_all_table_fonts(doc)
    enable_update_fields(doc)
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
