from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT = Path(__file__).resolve().parent
OUT = PROJECT / "10_output" / "ATMR多模态智能心理测评系统_毕业论文初稿_降AI率修订版.docx"


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_line: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    if first_line:
        fmt.first_line_indent = Pt(24)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 10, bold)


def set_cell_borders(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn(f"w:{key}"), str(value))


def apply_table_borders(table) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            borders = {
                "left": {"val": "nil"},
                "right": {"val": "nil"},
                "insideH": {"val": "nil"},
                "insideV": {"val": "nil"},
            }
            if row_index == 0:
                borders["top"] = {"val": "single", "sz": "12", "color": "000000"}
                borders["bottom"] = {"val": "single", "sz": "4", "color": "000000"}
            elif row_index == len(table.rows) - 1:
                borders["bottom"] = {"val": "single", "sz": "12", "color": "000000"}
            set_cell_borders(cell, **borders)


def clean_inline(text: str) -> str:
    text = re.sub(r"\[\[SYM:(.*?)\]\]", lambda m: math_text(m.group(1)), text)
    text = re.sub(r"\[\[REF_FIG:(.*?)\]\]", lambda m: "相关图", text)
    text = re.sub(r"\[\[REF_TBL:(.*?)\]\]", lambda m: "相关表", text)
    return text


def math_text(expr: str) -> str:
    expr = expr.strip()
    replacements = {
        r"\mu": "μ",
        r"\tau": "τ",
        r"\sum": "Σ",
        r"\frac": "",
        r"\hat": "",
        r"\in": "∈",
        r"\times": "×",
    }
    for key, value in replacements.items():
        expr = expr.replace(key, value)
    expr = expr.replace("{", "").replace("}", "")
    expr = expr.replace("_", "")
    expr = expr.replace("^", "")
    return expr


def equation_text(expr: str) -> str:
    expr = expr.strip()
    known = {
        r"\mu_t=\frac{\tau_{t-1}\mu_{t-1}+\tau_o x_t}{\tau_{t-1}+\tau_o}": "μt = (τt-1 μt-1 + τo xt) / (τt-1 + τo)",
        r"S(q)=w_1I(q)+w_2C(q)+w_3D(q)+w_4R(q)": "S(q)=w1I(q)+w2C(q)+w3D(q)+w4R(q)",
        r"MSE=\frac{1}{n}\sum_{i=1}^{n}(y_i-\hat{y_i})^2": "MSE = (1/n)Σ(yi - ŷi)^2",
        r"MAE=\frac{1}{n}\sum_{i=1}^{n}|y_i-\hat{y_i}|": "MAE = (1/n)Σ|yi - ŷi|",
    }
    return known.get(expr, math_text(expr))


def add_text_with_refs(paragraph, text: str) -> None:
    text = clean_inline(text)
    pos = 0
    for match in re.finditer(r"\[\[REF:(\d+)\]\]", text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, 12)
        run = paragraph.add_run(f"[{match.group(1)}]")
        set_run_font(run, 10)
        run.font.superscript = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, 12)


def add_normal_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=True)
    add_text_with_refs(p, text)


def add_center_paragraph(doc: Document, text: str, size: int = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p)
    run = p.add_run(text)
    set_run_font(run, size, bold)


def markdown_tables(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [cell.strip() for cell in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, caption: str, rows: list[list[str]], chapter_no: int, table_count: int) -> None:
    add_center_paragraph(doc, f"表 {chapter_no}-{table_count} {caption}", 11, False)
    if not rows:
        add_center_paragraph(doc, "（表格内容待补充）", 11)
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True
    for r, row in enumerate(rows):
        for c in range(col_count):
            set_cell_text(table.cell(r, c), row[c] if c < len(row) else "", bold=(r == 0))
    apply_table_borders(table)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True


def add_title_pages(doc: Document, state: dict) -> None:
    title = state["thesis"]["title"]
    add_center_paragraph(doc, "毕业设计（论文）", 22, True)
    doc.add_paragraph()
    add_center_paragraph(doc, title, 18, True)
    doc.add_paragraph()
    for label, value in [
        ("学生姓名", state["thesis"].get("student", "待补充")),
        ("专业", state["thesis"].get("major", "待补充")),
        ("学院", state["thesis"].get("school", "待补充")),
        ("指导教师", "待补充"),
        ("完成日期", "待补充"),
    ]:
        add_center_paragraph(doc, f"{label}：{value}", 12)
    doc.add_page_break()
    doc.add_heading("摘  要", level=1)
    add_normal_paragraph(doc, state["thesis"].get("abstractZh", "待补充"))
    add_normal_paragraph(doc, "关键词：ATMR；智能心理测评；多智能体辩论；RAG；多模态人格预测；AGTN-MTL")
    doc.add_page_break()
    doc.add_heading("Abstract", level=1)
    add_normal_paragraph(doc, state["thesis"].get("abstractEn", "To be completed."))
    add_normal_paragraph(doc, "Keywords: ATMR; intelligent psychological assessment; multi-agent debate; RAG; multimodal personality prediction; AGTN-MTL")
    doc.add_page_break()


def add_static_toc(doc: Document, chapter_files: list[Path]) -> None:
    doc.add_heading("目  录", level=1)
    for path in chapter_files:
        lines = path.read_text(encoding="utf-8").splitlines()
        for line in lines:
            if line.startswith("# "):
                p = doc.add_paragraph()
                set_paragraph_format(p)
                run = p.add_run(line[2:].strip())
                set_run_font(run, 12, True)
            elif line.startswith("## "):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                set_paragraph_format(p)
                run = p.add_run(line[3:].strip())
                set_run_font(run, 11)
    doc.add_page_break()


def add_chapter(doc: Document, path: Path, chapter_no: int) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    fig_count = 0
    table_count = 0
    eq_count = 0
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line or line.startswith("<!--"):
            i += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("[[FIG:"):
            fig_count += 1
            desc = line.removeprefix("[[FIG:").removesuffix("]]")
            add_center_paragraph(doc, f"（此处插入：{desc}）", 11)
            add_center_paragraph(doc, f"图 {chapter_no}-{fig_count} {desc}", 11)
        elif line.startswith("[[TBL:"):
            table_count += 1
            desc = line.removeprefix("[[TBL:").removesuffix("]]")
            rows, next_i = markdown_tables(lines, i + 1)
            add_table(doc, desc, rows, chapter_no, table_count)
            i = next_i - 1
        elif line.startswith("[[EQ:"):
            eq_count += 1
            expr = line.removeprefix("[[EQ:").removesuffix("]]")
            add_center_paragraph(doc, f"{equation_text(expr)}    ({chapter_no}.{eq_count})", 12)
        else:
            add_normal_paragraph(doc, line)
        i += 1


def add_references(doc: Document, references: list[dict]) -> None:
    doc.add_page_break()
    doc.add_heading("参考文献", level=1)
    for ref in references:
        p = doc.add_paragraph()
        set_paragraph_format(p)
        run = p.add_run(f"[{ref['id']}] {ref['description']}")
        set_run_font(run, 12)


def main() -> None:
    state = json.loads((PROJECT / "09_state" / "project_state.json").read_text(encoding="utf-8"))
    chapter_files = sorted((PROJECT / "03_chapters").glob("ch*_draft.md"))
    doc = Document()
    style_document(doc)
    add_title_pages(doc, state)
    add_static_toc(doc, chapter_files)
    for index, chapter_file in enumerate(chapter_files, start=1):
        if index > 1:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_chapter(doc, chapter_file, index)
    add_references(doc, state.get("references", []))
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
